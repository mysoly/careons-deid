"""
zorgdeid.processors.doc_processor
--------------------------------------
Document reading and PII processing pipelines for file-based input.

Reads .pdf, .docx, .txt, and image files (.jpg, .jpeg, .png, .tiff, .bmp,
.webp), extracts their plain text, then delegates to text_processor for
analysis and guarding.

OCR support
-----------
PDF files are first processed with pypdf (fast, zero-latency for digital
PDFs).  When a page yields fewer than OCR_FALLBACK_THRESHOLD characters
(default 50), it is treated as a scanned/image page and processed via EasyOCR
with the Dutch (``nl``) language model — **no system-level install required**.

Image files (.jpg, .jpeg, .png, .tiff, .bmp, .webp) are always processed
via OCR.

OCR requires the ``[ocr]`` optional dependencies only — no external binaries::

    pip install zorgdeid[ocr]

EasyOCR downloads the Dutch model (~120 MB) automatically on first use and
caches it in ``~/.EasyOCR/``.  PyMuPDF renders PDF pages as images with no
dependency on poppler or Ghostscript.

Consumed by the public namespace objects in zorgdeid/__init__.py::

    from zorgdeid import analyze, guard

    analyze.doc("/path/to/file.pdf")                     # -> list[dict]
    guard.doc("/path/to/file.docx", config={"mode": "tag"})  # -> dict
    guard.doc("/path/to/scan.png")                       # -> dict  (OCR)

Functions in this module
------------------------
    analyze(path, config) -> list[dict]   detect PII in a document file
    guard(path, config)   -> dict         detect + anonymize a document file
    read(path)            -> str          extract plain text only (no PII pipeline)

Exceptions
----------
    UnsupportedFormatError   raised for any extension not in SUPPORTED_EXTENSIONS
    FileNotFoundError        raised when the path does not exist (after the extension check)
    OcrNotAvailableError     raised when OCR is needed but easyocr/pymupdf are missing
"""

import logging
import os
import re
from typing import Dict, List, Optional, Set

from zorgdeid.processors.text_processor import analyze as _analyze, guard as _guard

logger = logging.getLogger(__name__)

# Minimum character count returned by pypdf before we attempt OCR for that page.
# Pages with fewer characters are assumed to be scanned/image-based.
OCR_FALLBACK_THRESHOLD: int = 50

# EasyOCR singleton — loaded once on first OCR call, then reused.
# Loading takes ~2–5 s; subsequent calls are fast.
_easyocr_reader = None

SUPPORTED_EXTENSIONS: Set[str] = {
    ".pdf", ".docx", ".txt",
    # image formats (always OCR)
    ".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp",
}

IMAGE_EXTENSIONS: Set[str] = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp"}


class OcrNotAvailableError(ImportError):
    """Raised when OCR is required but ``easyocr`` / ``pymupdf`` are missing."""


class UnsupportedFormatError(ValueError):
    """Raised when a file with an unsupported extension is provided."""


# ---------------------------------------------------------------------------
# File reader
# ---------------------------------------------------------------------------

def read(path: str) -> str:
    """
    Extract plain text from a document file.

    Parameters
    ----------
    path : str
        Absolute or relative path to the file.
        Supported extensions: ``.pdf``, ``.docx``, ``.txt``,
        ``.jpg``, ``.jpeg``, ``.png``, ``.tiff``, ``.tif``, ``.bmp``, ``.webp``

    Returns
    -------
    str
        Plain text content of the document.

    Raises
    ------
    UnsupportedFormatError
        If the file extension is not in SUPPORTED_EXTENSIONS.
    FileNotFoundError
        If the file does not exist.
    OcrNotAvailableError
        If the file requires OCR but ``easyocr`` / ``pymupdf`` are not
        installed.
    """
    ext = os.path.splitext(path)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFormatError(
            f"Unsupported format {ext!r}. Supported formats are: {supported}"
        )

    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path!r}")

    if ext == ".txt":
        return _read_txt(path)
    elif ext == ".pdf":
        return _read_pdf(path)
    elif ext in IMAGE_EXTENSIONS:
        return _read_image(path)
    else:
        return _read_docx(path)


# ---------------------------------------------------------------------------
# Processing pipelines
# ---------------------------------------------------------------------------

def analyze(path: str, config: Optional[Dict] = None) -> List[Dict]:
    """
    Read a document and return a list of PII findings.

    Parameters
    ----------
    path   : Path to .pdf, .docx, or .txt file.
    config : Optional detection config (same as text_processor.analyze).

    Returns
    -------
    list[dict]
        Each dict: ``{"type": str, "start": int, "end": int, "score": float}``

    Raises
    ------
    UnsupportedFormatError
        If the file extension is not .pdf, .docx, or .txt.
    FileNotFoundError
        If *path* does not exist.
    """
    return _analyze(read(path), config=config)


def guard(path: str, config: Optional[Dict] = None) -> Dict:
    """
    Read a document, anonymize its PII, and return the guarded result.

    Parameters
    ----------
    path   : Path to .pdf, .docx, or .txt file.
    config : Optional processing config (same as text_processor.guard).

    Returns
    -------
    dict
        ``guarded_text`` – text with PII replaced.
        ``findings``     – list of finding dicts.

    Raises
    ------
    UnsupportedFormatError
        If the file extension is not .pdf, .docx, or .txt.
    FileNotFoundError
        If *path* does not exist.
    """
    return _guard(read(path), config=config)


# ---------------------------------------------------------------------------
# Format-specific readers
# ---------------------------------------------------------------------------

def _read_txt(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8") as fh:
            return fh.read()
    except UnicodeDecodeError:
        with open(path, "r", encoding="cp1252", errors="replace") as fh:
            return fh.read()
    except PermissionError:
        raise PermissionError(f"Permission denied: {path!r}")
    except IsADirectoryError:
        raise IsADirectoryError(f"Is a directory, not a file: {path!r}")


def _normalize_pdf_text(text: str) -> str:
    """
    Repair the two most common pypdf extraction artefacts:

    1. Word-per-line scattering — when pypdf emits each word on its own line
       separated by a whitespace-only line:
           ``aanhoudende\\n \\npijn\\n \\nop``
       These are collapsed back into a single space.

    2. Double (or triple) spacing — pypdf inserts explicit character spacing
       that results in ``Medisch  Verslag:`` instead of ``Medisch Verslag:``.
       All runs of 2+ spaces/tabs are reduced to one space.

    Genuine paragraph breaks (blank lines) are preserved up to one blank line.
    """

    text = re.sub(r"\n[ \t]+\n", " ", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return text.strip()


def _get_easyocr_reader():
    """
    Return the shared EasyOCR Reader for Dutch text (singleton).

    The reader is initialised once on first call (~2–5 s; downloads the Dutch
    model ~120 MB to ``~/.EasyOCR/`` if not already cached) and then reused
    for all subsequent OCR calls in the same process.

    Raises ``OcrNotAvailableError`` if ``easyocr`` is not installed.
    """
    global _easyocr_reader
    if _easyocr_reader is not None:
        return _easyocr_reader

    try:
        import easyocr
    except ImportError as exc:
        raise OcrNotAvailableError(
            "easyocr is required for OCR support. "
            "Install the OCR extras with:  pip install zorgdeid[ocr]\n"
            "The Dutch language model (~120 MB) will be downloaded automatically "
            "on first use and cached in ~/.EasyOCR/."
        ) from exc

    logger.info(
        "Initialising EasyOCR Dutch model — "
        "first run downloads ~120 MB to ~/.EasyOCR/ (cached afterwards)."
    )
    _easyocr_reader = easyocr.Reader(["nl"], gpu=False, verbose=False)
    return _easyocr_reader


def _ocr_image(image: "PIL.Image.Image") -> str:  # type: ignore[name-defined]  # noqa: F821
    """
    Run EasyOCR on a single PIL Image and return extracted text.

    Uses the Dutch (``nl``) language model.  No system-level binary required.
    Raises ``OcrNotAvailableError`` if easyocr is not installed.
    """
    import io

    reader = _get_easyocr_reader()

    buf = io.BytesIO()
    image.save(buf, format="PNG")
    raw_bytes = buf.getvalue()

    lines = reader.readtext(raw_bytes, detail=0, paragraph=True)
    return "\n".join(lines)


def _read_pdf(path: str) -> str:
    """
    Extract text from a PDF file.

    Digital PDFs are processed with pypdf (fast, no OCR).  Pages that yield
    fewer than ``OCR_FALLBACK_THRESHOLD`` characters are assumed to be scanned
    and re-processed via EasyOCR (Dutch model, no system install required).
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError(
            "pypdf is required for PDF support. "
            "Install it with:  pip install pypdf"
        ) from exc

    reader = PdfReader(path)
    pages_text: List[str] = []

    for page_num, page in enumerate(reader.pages):
        text = page.extract_text() or ""

        if len(text.strip()) < OCR_FALLBACK_THRESHOLD:
            logger.debug(
                "PDF page %d yielded %d chars — falling back to EasyOCR.",
                page_num + 1,
                len(text.strip()),
            )
            text = _ocr_pdf_page(path, page_num)

        pages_text.append(text)

    return _normalize_pdf_text("\n".join(pages_text))


def _ocr_pdf_page(path: str, page_num: int) -> str:
    """
    Render a single PDF page as an image and run EasyOCR on it.

    Uses PyMuPDF (``fitz``) for high-fidelity rendering at 2× zoom (~144 DPI).
    No poppler or Ghostscript dependency.
    Raises ``OcrNotAvailableError`` if pymupdf or easyocr are missing.
    """
    try:
        import fitz  # pymupdf
    except ImportError as exc:
        raise OcrNotAvailableError(
            "pymupdf is required for scanned PDF OCR support. "
            "Install the OCR extras with:  pip install zorgdeid[ocr]"
        ) from exc

    try:
        from PIL import Image
        import io
    except ImportError as exc:
        raise OcrNotAvailableError(
            "Pillow is required for image OCR support. "
            "Install the OCR extras with:  pip install zorgdeid[ocr]"
        ) from exc

    doc = fitz.open(path)
    page = doc[page_num]
    # 2× zoom ≈ 144 DPI — good balance of OCR accuracy and speed
    mat = fitz.Matrix(2, 2)
    pix = page.get_pixmap(matrix=mat)
    doc.close()

    img = Image.open(io.BytesIO(pix.tobytes("png")))
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    return _ocr_image(img)


def _read_image(path: str) -> str:
    """
    Read a raster image file and extract text via EasyOCR (Dutch model).

    Supported formats: JPEG, PNG, TIFF, BMP, WebP.
    Raises ``OcrNotAvailableError`` if easyocr or Pillow are missing.
    """
    try:
        from PIL import Image
    except ImportError as exc:
        raise OcrNotAvailableError(
            "Pillow is required for image OCR support. "
            "Install the OCR extras with:  pip install zorgdeid[ocr]"
        ) from exc

    with Image.open(path) as img:
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        text = _ocr_image(img)

    return text.strip()


def _read_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX support. "
            "Install it with:  pip install python-docx"
        ) from exc

    doc = Document(path)
    parts: List[str] = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)
